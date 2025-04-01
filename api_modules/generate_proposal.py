# Standard library imports
import asyncio
import io
import re
import time
from typing import List

# Third-party library imports
from bson import ObjectId, Binary
from docx import Document
from fastapi import HTTPException

# Local/project-specific imports
from api_modules.tender_library import (
    log,
)
from config import bids_collection
from services.chain import invoke_graph
from utils import calculate_word_amounts, has_permission_to_access_bid

async def process_generate_proposal(
    bid_id: str, selected_folders: List[str], current_user: str
) -> dict:
    try:
        log.info(f"Starting proposal generation for bid {bid_id}")
        start_time = time.time()
        # Run document setup and bid fetch concurrently
        bid = await bids_collection.find_one({"_id": ObjectId(bid_id)})
        if not bid:
            raise HTTPException(status_code=404, detail="Bid not found")
        if not await has_permission_to_access_bid(bid, current_user):
            raise HTTPException(
                status_code=403, detail="You don't have permission to access this bid"
            )
        outline = bid.get("outline", [])
        if not outline:
            raise HTTPException(status_code=404, detail="No outline sections found")

        selected_case_studies = bid.get("selectedCaseStudies", [])
        selected_case_studies_raw_text = []

        for document in selected_case_studies:
            if "rawtext" in document:
                selected_case_studies_raw_text.append(document["rawtext"])
            else:
                log.warning(f"No rawtext field found for document: {document.get('name', 'unknown')}")

        solution = bid.get("solution", {})

        if "product" in solution:
            solution["product"] = "What specific product or service are you bidding to provide? " + solution["product"]

        if "features" in solution:
            solution["features"] = "What are some key technical features or capabilities of your solution? " + solution["features"]

        if "approach" in solution:
            solution["approach"] = "What makes your approach technically superior to alternatives? " + solution["approach"]

        # Process sections in batches to avoid overwhelming the LLM service
        BATCH_SIZE = 50  # Adjust based on your LLM service limits
        sections = []
        updated_outline = []  # Store the updated outline with answers

        for i in range(0, len(outline), BATCH_SIZE):
            batch = outline[i : i + BATCH_SIZE]
            batch_tasks = [
                process_section(section, selected_folders, current_user, bid_id, selected_case_studies_raw_text, solution)
                for section in batch
            ]
            batch_results = await asyncio.gather(*batch_tasks)
            sections.extend(batch_results)

            # Update the outline entries with their answers
            for original_section, result in zip(batch, batch_results):
                section_copy = original_section.copy()
                section_copy["answer"] = result["answer"]
                updated_outline.append(section_copy)

        # Update the entire outline in one operation
        await bids_collection.update_one(
            {"_id": ObjectId(bid_id)}, {"$set": {"outline": updated_outline}}
        )

        total_time = time.time() - start_time
        log.info(f"Total processing time: {total_time:.2f} seconds")
        return {
            "updated_outline": updated_outline,
        }
    except Exception as e:
        log.error(f"Fatal error in process_generate_proposal: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


async def process_section(section, selected_folders, current_user, bid_id, selected_case_studies_raw_text, solution):
    try:
        """Process a single section with its LLM call"""
        section_start = time.time()
        section_title = section["heading"].replace("|", "").strip()
        question = section["question"]

        subsection_compliance = section.get("compliance_requirements", " ")
        if subsection_compliance is None:
            log.warning("No compliance requirements found for this section.")

        subsection_evaluation_criteria = section.get("relevant_evaluation_criteria", " ")
        if subsection_evaluation_criteria is None:
            log.warning("No evaluation criteria found for this section.")

        subsection_derived_insights = section.get("relevant_derived_insights", " ")
        if subsection_derived_insights is None:
            log.warning("No derived insights found for this section.")

        subsection_differentiation_factors = section.get("relevant_differentiation_factors", " ")
        if subsection_differentiation_factors is None:
            log.warning("No differentiation factors found for this section.")

        highlighted_documents_raw_text = []
        highlighted_documents = section.get("highlightedDocuments", [])

        for document in highlighted_documents:
            # Extract the rawtext field from each document
            if "rawtext" in document:
                # Print the first 20 characters of the rawtext
                log.info(f"First 20 chars of {document.get('name', 'unknown')}: {document['rawtext'][:20]}")
                highlighted_documents_raw_text.append(document["rawtext"])
            else:
                log.warning(f"No rawtext field found for document: {document.get('name', 'unknown')}")

        log.info("highlighted_documents")
        log.info(len(highlighted_documents))
        log.info(len(highlighted_documents_raw_text))

        log.info(f"Starting section: {question}")
        try:
            word_amounts = calculate_word_amounts(section["subheadings"], section["word_count"])
        except Exception as e:
            log.error(f"Error calculating word amounts: {str(e)} for section {section_title} of bid {bid_id} user {current_user}")
            word_amounts = [500]
        subtopics = []
        writingplans = []
        compliance_reqs = []
        evaluation_criteria = []
        derived_insights = []
        differentation_factors = []

        if "subheadings" in section:
            for subheading in section["subheadings"]:
                subtopics.append(subheading["title"])
                compliance_reqs.append(subsection_compliance)
                evaluation_criteria.append(subsection_evaluation_criteria)
                derived_insights.append(subsection_derived_insights)
                differentation_factors.append(subsection_differentiation_factors)
                writingplans.append(subheading["extra_instructions"])

        if len(subtopics) == 0:
            subtopics.append(section_title)
            compliance_reqs.append(subsection_compliance)
            evaluation_criteria.append(subsection_evaluation_criteria)
            derived_insights.append(subsection_derived_insights)
            differentation_factors.append(subsection_differentiation_factors)
            writingplans.append("")

        choice = section.get("choice", "3b")
        max_retries = 3
        base_delay = 5
        backgroundinfo = section.get("writingplan", "") #we are not using this anymore

        log.info(question)
        log.info(word_amounts)
        log.info(subtopics)
        log.info(choice)
        log.info(backgroundinfo)

        for attempt in range(max_retries):
            try:
                answer = await invoke_graph(
                    choice,
                    question,
                    backgroundinfo,
                    current_user,
                    selected_folders,
                    "4",
                    bid_id,
                    subtopics,
                    word_amounts,
                    compliance_reqs,
                    evaluation_criteria,
                    derived_insights,
                    differentation_factors,
                    writingplans,
                    highlighted_documents_raw_text,
                    selected_case_studies_raw_text,
                    solution
                )

                log.info(f"Completed section: {question} in {time.time() - section_start:.2f} seconds")
                return {"answer": answer, "question": question, "section_title": section_title}

            except Exception as e:
                if attempt < max_retries - 1:
                    delay = base_delay * (2**attempt)
                    log.warning(f"API overloaded, retrying in {delay} seconds...")
                    await asyncio.sleep(delay)
                    continue

                log.error(f"Error processing section: {question} - {str(e)}", exc_info=True)
                raise
    except Exception as e:
        log.error(f"Critical error in process_section: {str(e)}", exc_info=True)
        return {
            "answer": f"Critical error processing section: {str(e)}",
            "question": section.get("question", "Unknown question"),
            "section_title": section.get("heading", "Unknown section").replace("|", "").strip(),
        }

async def remove_references(bid_id: str, current_user: str) -> dict:
    """
    Removes all references in square brackets from the generated proposal document.
    """
    try:
        # Fetch the bid document
        bid = await bids_collection.find_one({"_id": ObjectId(bid_id)})
        if not bid:
            raise HTTPException(status_code=404, detail="Bid not found")
        if not await has_permission_to_access_bid(bid, current_user):
            raise HTTPException(
                status_code=403, detail="You don't have permission to access this bid"
            )

        # Get the generated proposal
        doc_content = bid.get("generated_proposal")
        if not doc_content:
            raise HTTPException(status_code=404, detail="No generated proposal found")

        # Load the document from bytes
        doc_buffer = io.BytesIO(doc_content)
        doc = Document(doc_buffer)

        # Remove references from each paragraph using regex
        for paragraph in doc.paragraphs:
            # Remove text within square brackets
            cleaned_text = re.sub(r'\[.*?\]', '', paragraph.text)
            if cleaned_text != paragraph.text:  # Only modify if changes were made
                paragraph.clear()
                paragraph.add_run(cleaned_text)

        # Save the modified document
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        doc_content = output_buffer.getvalue()

        # Update the document in MongoDB
        await bids_collection.update_one(
            {"_id": ObjectId(bid_id)}, {"$set": {"generated_proposal": Binary(doc_content)}}
        )

        return {
            "filename": f"proposal_{bid_id}.docx",
            "content": doc_content,
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    except Exception as e:
        log.error(f"Error removing references: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
