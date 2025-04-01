from enum import IntEnum
import logging
import re
import io
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from bson.objectid import ObjectId
from config import bids_collection

log = logging.getLogger(__name__)

# Define Word specific highlight colors for reference
class WD_HIGHLIGHT_COLOR(IntEnum):
    YELLOW = 1
    GREEN = 2
    BLUE = 3
    RED = 4
    # Add more as needed

def process_html_content(html_content):
    """
    Process HTML content and prepare it for insertion into a Word document
    Returns a list of instructions for creating paragraphs with proper formatting
    """
    if not html_content:
        return []

    # Check if content is actually HTML
    if not any(tag in html_content for tag in ['<p', '<div', '<span', '<h1', '<h2', '<h3', '<strong', '<em', '<ul', '<ol']):
        # Handle as plain text with custom markers
        return process_section_content(html_content)

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process all top-level elements in the body
    elements = []
    # Get the body or use the whole soup if no body tag
    body = soup.body if soup.body else soup

    for element in body.find_all(recursive=False):
        if element.name in ['p', 'div']:
            elements.append(('paragraph', extract_formatting(element)))
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            elements.append(('heading', element.get_text().strip(), level))
        elif element.name == 'ul':
            elements.append(('list', extract_list_items(element), 'bullet'))
        elif element.name == 'ol':
            elements.append(('list', extract_list_items(element), 'number'))
        elif element.name == 'blockquote':
            elements.append(('quote', extract_formatting(element)))
        else:
            # Handle any other elements as plain paragraphs
            elements.append(('paragraph', extract_formatting(element)))

    # If no elements were found, try to extract text directly
    if not elements and body.get_text().strip():
        elements.append(('paragraph', [(body.get_text().strip(), {})]))

    return elements

def extract_formatting(element):
    """
    Extract text with formatting instructions from an HTML element
    Returns a list of (text, formatting) tuples
    """
    result = []

    # Process children with their formatting
    for child in element.children:
        if isinstance(child, str):
            # Plain text
            text = child
            if text.strip():
                result.append((text, {}))
        elif child.name in ('strong', 'b'):
            # Bold text
            text = child.get_text()
            if text.strip():
                result.append((text, {'bold': True}))
        elif child.name in ('em', 'i'):
            # Italic text
            text = child.get_text()
            if text.strip():
                result.append((text, {'italic': True}))
        elif child.name == 'u':
            # Underlined text
            text = child.get_text()
            if text.strip():
                result.append((text, {'underline': True}))
        elif child.name == 'a':
            # Hyperlink
            text = child.get_text()
            href = child.get('href', '')
            if text.strip():
                result.append((text, {'hyperlink': href}))
        elif child.name == 'br':
            # Line break - add empty string with line break flag
            result.append(('', {'break': True}))
        elif child.name == 'span':
            # Process span and its children recursively
            nested_results = extract_formatting(child)

            # Check for background color in style attribute
            style = child.get('style', '')
            bg_color = None
            if 'background-color' in style:
                match = re.search(r'background-color:\s*([^;]+)', style)
                if match:
                    bg_color = match.group(1).strip()

            # Apply background color to all nested elements if it exists
            if bg_color:
                for i, (text, formatting) in enumerate(nested_results):
                    formatting['highlight'] = bg_color
                    nested_results[i] = (text, formatting)

            result.extend(nested_results)
        else:
            # For any other elements, just get the text
            text = child.get_text()
            if text.strip():
                result.append((text, {}))

    return result

def extract_list_items(list_element):
    """
    Extract items from ul/ol elements
    Returns a list of formatted items
    """
    items = []
    for li in list_element.find_all('li', recursive=False):
        items.append(extract_formatting(li))
    return items

def apply_formatting_to_docx(doc, elements):
    """
    Apply the extracted formatting instructions to the Word document
    """
    for element_type, *args in elements:
        if element_type == 'paragraph':
            formatting_instructions = args[0]
            paragraph = doc.add_paragraph()

            for text, formatting in formatting_instructions:
                if not text and formatting.get('break'):
                    # Handle line break
                    paragraph.add_run().add_break()
                    continue

                run = paragraph.add_run(text)

                # Apply formatting
                if formatting.get('bold'):
                    run.bold = True
                if formatting.get('italic'):
                    run.italic = True
                if formatting.get('underline'):
                    run.underline = True
                if formatting.get('hyperlink'):
                    # In a real implementation, you'd use the proper Word API for hyperlinks
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
                if formatting.get('highlight'):
                    # Convert HTML color to Word highlight color (simplified)
                    # In production, you would need a proper color mapping
                    color = formatting.get('highlight').lower()

                    # Simple color mapping example
                    if 'yellow' in color or '#ffff' in color:
                        run.font.highlight_color = WD_HIGHLIGHT_COLOR.YELLOW
                    elif 'green' in color or '#00ff' in color:
                        run.font.highlight_color = WD_HIGHLIGHT_COLOR.GREEN
                    elif 'blue' in color or '#0000ff' in color:
                        run.font.highlight_color = WD_HIGHLIGHT_COLOR.BLUE
                    elif 'red' in color or '#ff00' in color:
                        run.font.highlight_color = WD_HIGHLIGHT_COLOR.RED
                    else:
                        # Default to yellow for any other color
                        run.font.highlight_color = WD_HIGHLIGHT_COLOR.YELLOW

        elif element_type == 'heading':
            text, level = args
            doc.add_heading(text, level=level)

        elif element_type == 'list':
            items, list_type = args
            for item_formatting in items:
                if list_type == 'bullet':
                    paragraph = doc.add_paragraph(style='List Bullet')
                else:
                    paragraph = doc.add_paragraph(style='List Number')

                for text, formatting in item_formatting:
                    run = paragraph.add_run(text)
                    if formatting.get('bold'):
                        run.bold = True
                    if formatting.get('italic'):
                        run.italic = True
                    if formatting.get('underline'):
                        run.underline = True

        elif element_type == 'quote':
            formatting_instructions = args[0]
            paragraph = doc.add_paragraph(style='Intense Quote')

            for text, formatting in formatting_instructions:
                run = paragraph.add_run(text)
                if formatting.get('bold'):
                    run.bold = True
                if formatting.get('italic'):
                    run.italic = True

def process_section_content(text):
    """
    Process plain text content with custom markers
    This returns element instructions for non-HTML content
    """
    elements = []
    segments = []
    current_position = 0

    for match in re.finditer(r'###(.*?)###\s*:', text):
        if match.start() > current_position:
            content = text[current_position : match.start()].strip()
            if content:
                segments.append(('normal', content))
        subheading = match.group(1).strip()
        if subheading:
            segments.append(('subheading', subheading))
        current_position = match.end()

    if current_position < len(text):
        content = text[current_position:].strip()
        if content:
            segments.append(('normal', content))

    # Convert segments to elements format
    for segment_type, content in segments:
        if segment_type == 'subheading':
            elements.append(('heading', content, 3))
        else:
            elements.append(('paragraph', [
                (content, {})
            ]))

    # If no segments were processed but text exists, add it as a paragraph
    if not segments and text.strip():
        elements.append(('paragraph', [
            (text.strip(), {})
        ]))

    return elements

async def process_generate_docx(bid_id: str) -> StreamingResponse:
    """
    Generate a DOCX file from bid outline sections with proper HTML handling
    """
    try:
        log.info(f"Starting document generation for bid {bid_id}")
        start_time = time.time()

        # Fetch the bid from the database
        bid = await bids_collection.find_one({"_id": ObjectId(bid_id)})
        if not bid:
            raise HTTPException(status_code=404, detail="Bid not found")

        # Get the outline sections
        outline = bid.get("outline", [])
        if not outline:
            raise HTTPException(status_code=404, detail="No outline sections found")

        # Get the bid name/title
        bid_info = bid.get("bidInfo", "Proposal Document")

        # Create a new document
        doc = Document()

        # Set up document styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Configure title style
        title_style = doc.styles['Title']
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.name = 'Calibri'

        # Add main title
        title = doc.add_paragraph()
        title.style = title_style
        title.add_run(bid_info)
        title.alignment = 1  # Center alignment

        # Add each section to the document
        for section in outline:
            # Add section heading
            doc.add_heading(section.get("heading", "Untitled Section"), level=1)

            # Add section question if it exists
            if section.get("question"):
                question_para = doc.add_paragraph()
                question_para.style = "Intense Quote"
                question_run = question_para.add_run("Question: ")
                question_run.bold = True
                question_para.add_run(section["question"])

            # Add section content - this is where we now handle HTML
            if section.get("answer"):
                # Process the section content with HTML handling
                elements = process_html_content(section["answer"])

                # Apply the formatted elements to the document
                apply_formatting_to_docx(doc, elements)

            # Add space after section
            doc.add_paragraph()

        # Save to memory
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        log.info(f"Document generation completed in {time.time() - start_time:.2f} seconds")

        # Return the document as a streaming response
        filename = f"proposal_{bid_id}.docx"
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        return StreamingResponse(docx_io, headers=headers)

    except Exception as e:
        log.error(f"Error in process_generate_docx: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
